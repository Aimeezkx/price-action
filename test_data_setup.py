#!/usr/bin/env python3
"""
Test data setup script for comprehensive testing.
Creates sample documents, test databases, and mock data.
"""
import asyncio
import json
import os
import shutil
from pathlib import Path
from typing import Dict, List

import click
from faker import Faker
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document as DocxDocument
from docx.shared import Inches

fake = Faker()


class TestDataGenerator:
    """Generate comprehensive test data for the application."""
    
    def __init__(self, output_dir: Path = Path("test_data")):
        self.output_dir = output_dir
        self.output_dir.mkdir(exist_ok=True)
        
    def create_sample_pdf(self, filename: str, pages: int = 10, with_images: bool = True) -> Path:
        """Create a sample PDF document."""
        pdf_path = self.output_dir / filename
        
        doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title page
        title = Paragraph("Sample Document for Testing", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Table of contents
        toc_title = Paragraph("Table of Contents", styles['Heading1'])
        story.append(toc_title)
        
        for i in range(1, pages):
            chapter_title = f"Chapter {i}: {fake.sentence(nb_words=4)}"
            toc_entry = Paragraph(f"{chapter_title} ........................ {i+1}", styles['Normal'])
            story.append(toc_entry)
        
        story.append(Spacer(1, 24))
        
        # Content pages
        for i in range(1, pages):
            # Chapter heading
            chapter_title = f"Chapter {i}: {fake.sentence(nb_words=4)}"
            heading = Paragraph(chapter_title, styles['Heading1'])
            story.append(heading)
            story.append(Spacer(1, 12))
            
            # Chapter content
            for _ in range(3):
                # Section heading
                section_title = fake.sentence(nb_words=3)
                section_heading = Paragraph(section_title, styles['Heading2'])
                story.append(section_heading)
                story.append(Spacer(1, 6))
                
                # Paragraphs
                for _ in range(2):
                    paragraph_text = fake.paragraph(nb_sentences=5)
                    paragraph = Paragraph(paragraph_text, styles['Normal'])
                    story.append(paragraph)
                    story.append(Spacer(1, 6))
                
                # Add definition
                term = fake.word().title()
                definition = f"{term} is defined as {fake.sentence(nb_words=10)}"
                definition_para = Paragraph(f"<b>Definition:</b> {definition}", styles['Normal'])
                story.append(definition_para)
                story.append(Spacer(1, 12))
            
            # Add image placeholder if requested
            if with_images and i % 2 == 0:
                # Create a simple colored rectangle as image placeholder
                img_path = self.create_placeholder_image(f"figure_{i}.png")
                if img_path.exists():
                    img = Image(str(img_path), width=4*Inches, height=2*Inches)
                    story.append(img)
                    
                    caption = f"Figure {i//2}: {fake.sentence(nb_words=6)}"
                    caption_para = Paragraph(caption, styles['Caption'])
                    story.append(caption_para)
                    story.append(Spacer(1, 12))
        
        doc.build(story)
        return pdf_path
    
    def create_placeholder_image(self, filename: str) -> Path:
        """Create a placeholder image for testing."""
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            img_path = self.output_dir / filename
            
            # Create a simple colored rectangle
            img = Image.new('RGB', (400, 200), color='lightblue')
            draw = ImageDraw.Draw(img)
            
            # Add some text
            try:
                font = ImageFont.load_default()
            except:
                font = None
            
            text = f"Test Image\n{filename}"
            draw.text((50, 80), text, fill='black', font=font)
            
            # Add some shapes
            draw.rectangle([50, 50, 150, 100], outline='red', width=2)
            draw.ellipse([200, 50, 300, 150], outline='green', width=2)
            
            img.save(img_path)
            return img_path
            
        except ImportError:
            # If PIL is not available, create empty file
            img_path = self.output_dir / filename
            img_path.touch()
            return img_path
    
    def create_sample_docx(self, filename: str, pages: int = 5) -> Path:
        """Create a sample DOCX document."""
        docx_path = self.output_dir / filename
        
        doc = DocxDocument()
        
        # Title
        title = doc.add_heading('Sample DOCX Document', 0)
        
        # Table of contents
        doc.add_heading('Table of Contents', level=1)
        for i in range(1, pages):
            doc.add_paragraph(f"Chapter {i}: {fake.sentence(nb_words=4)}")
        
        # Content
        for i in range(1, pages):
            doc.add_heading(f'Chapter {i}: {fake.sentence(nb_words=4)}', level=1)
            
            for _ in range(2):
                doc.add_heading(fake.sentence(nb_words=3), level=2)
                
                for _ in range(3):
                    doc.add_paragraph(fake.paragraph(nb_sentences=4))
                
                # Add definition
                term = fake.word().title()
                definition = f"{term}: {fake.sentence(nb_words=8)}"
                p = doc.add_paragraph()
                p.add_run('Definition: ').bold = True
                p.add_run(definition)
        
        doc.save(docx_path)
        return docx_path
    
    def create_sample_markdown(self, filename: str) -> Path:
        """Create a sample Markdown document."""
        md_path = self.output_dir / filename
        
        content = f"""# Sample Markdown Document

This is a test document for the document learning application.

## Table of Contents

"""
        
        # Generate chapters
        chapters = []
        for i in range(1, 6):
            chapter_title = fake.sentence(nb_words=4)
            chapters.append((i, chapter_title))
            content += f"- [Chapter {i}: {chapter_title}](#chapter-{i})\n"
        
        content += "\n"
        
        # Generate content
        for i, chapter_title in chapters:
            content += f"## Chapter {i}: {chapter_title}\n\n"
            
            for j in range(1, 4):
                section_title = fake.sentence(nb_words=3)
                content += f"### {section_title}\n\n"
                
                # Add paragraphs
                for _ in range(2):
                    content += fake.paragraph(nb_sentences=4) + "\n\n"
                
                # Add definition
                term = fake.word().title()
                definition = fake.sentence(nb_words=10)
                content += f"**Definition**: {term} is defined as {definition}\n\n"
                
                # Add code block occasionally
                if j == 2:
                    content += "```python\n"
                    content += f"def {fake.word()}():\n"
                    content += f'    return "{fake.sentence()}"\n'
                    content += "```\n\n"
                
                # Add image reference
                if i % 2 == 0 and j == 1:
                    img_filename = f"figure_{i}.png"
                    self.create_placeholder_image(img_filename)
                    content += f"![Figure {i}]({img_filename})\n\n"
                    content += f"*Figure {i}: {fake.sentence(nb_words=6)}*\n\n"
        
        md_path.write_text(content)
        return md_path
    
    def create_test_database_data(self) -> Dict:
        """Create test data for database seeding."""
        data = {
            "documents": [],
            "chapters": [],
            "knowledge_points": [],
            "cards": [],
            "srs_states": []
        }
        
        # Create documents
        for i in range(5):
            doc_id = fake.uuid4()
            document = {
                "id": doc_id,
                "filename": f"test_document_{i+1}.pdf",
                "file_type": "pdf",
                "status": "completed",
                "metadata": {
                    "pages": fake.random_int(min=5, max=50),
                    "size": fake.random_int(min=1000, max=10000000),
                    "processing_time": fake.random_int(min=10, max=300)
                },
                "created_at": fake.date_time_this_year().isoformat()
            }
            data["documents"].append(document)
            
            # Create chapters for each document
            num_chapters = fake.random_int(min=2, max=8)
            for j in range(num_chapters):
                chapter_id = fake.uuid4()
                chapter = {
                    "id": chapter_id,
                    "document_id": doc_id,
                    "title": f"Chapter {j+1}: {fake.sentence(nb_words=4)}",
                    "level": 1,
                    "order_index": j + 1,
                    "page_start": j * 5 + 1,
                    "page_end": (j + 1) * 5
                }
                data["chapters"].append(chapter)
                
                # Create knowledge points for each chapter
                num_knowledge = fake.random_int(min=5, max=15)
                for k in range(num_knowledge):
                    knowledge_id = fake.uuid4()
                    knowledge_types = ["definition", "fact", "theorem", "process", "example"]
                    knowledge = {
                        "id": knowledge_id,
                        "chapter_id": chapter_id,
                        "kind": fake.random_element(knowledge_types),
                        "text": fake.paragraph(nb_sentences=3),
                        "entities": [fake.word() for _ in range(fake.random_int(min=1, max=5))],
                        "anchors": {
                            "page": fake.random_int(min=chapter["page_start"], max=chapter["page_end"]),
                            "chapter": j + 1,
                            "position": fake.random_int(min=0, max=1000)
                        },
                        "embedding": [fake.pyfloat(min_value=-1, max_value=1) for _ in range(384)]
                    }
                    data["knowledge_points"].append(knowledge)
                    
                    # Create cards for knowledge points
                    card_types = ["qa", "cloze", "image_hotspot"]
                    card_type = fake.random_element(card_types)
                    
                    card_id = fake.uuid4()
                    card = {
                        "id": card_id,
                        "knowledge_id": knowledge_id,
                        "card_type": card_type,
                        "front": self.generate_card_front(knowledge, card_type),
                        "back": self.generate_card_back(knowledge, card_type),
                        "difficulty": fake.pyfloat(min_value=0.1, max_value=1.0),
                        "metadata": self.generate_card_metadata(card_type)
                    }
                    data["cards"].append(card)
                    
                    # Create SRS state for card
                    srs = {
                        "id": fake.uuid4(),
                        "card_id": card_id,
                        "user_id": fake.uuid4(),
                        "ease_factor": fake.pyfloat(min_value=1.3, max_value=3.0),
                        "interval": fake.random_int(min=1, max=365),
                        "repetitions": fake.random_int(min=0, max=10),
                        "due_date": fake.date_time_between(start_date='-30d', end_date='+30d').isoformat(),
                        "last_reviewed": fake.date_time_this_month().isoformat()
                    }
                    data["srs_states"].append(srs)
        
        return data
    
    def generate_card_front(self, knowledge: Dict, card_type: str) -> str:
        """Generate card front based on type."""
        if card_type == "qa":
            if knowledge["kind"] == "definition" and knowledge["entities"]:
                return f"What is {knowledge['entities'][0]}?"
            else:
                return f"What do you know about {fake.word()}?"
        elif card_type == "cloze":
            text = knowledge["text"]
            entities = knowledge["entities"][:2]  # Max 2 blanks
            for entity in entities:
                text = text.replace(entity, "___", 1)
            return text
        elif card_type == "image_hotspot":
            return f"test_images/figure_{fake.random_int(min=1, max=10)}.png"
        
        return knowledge["text"][:100] + "..."
    
    def generate_card_back(self, knowledge: Dict, card_type: str) -> str:
        """Generate card back based on type."""
        if card_type == "image_hotspot":
            return f"Image caption: {fake.sentence(nb_words=8)}"
        return knowledge["text"]
    
    def generate_card_metadata(self, card_type: str) -> Dict:
        """Generate card metadata based on type."""
        if card_type == "cloze":
            return {
                "blanks": [fake.word() for _ in range(fake.random_int(min=1, max=3))],
                "original_text": fake.sentence()
            }
        elif card_type == "image_hotspot":
            return {
                "hotspots": [
                    {
                        "x": fake.random_int(min=50, max=300),
                        "y": fake.random_int(min=50, max=200),
                        "width": fake.random_int(min=30, max=80),
                        "height": fake.random_int(min=20, max=60),
                        "label": fake.word().title()
                    }
                    for _ in range(fake.random_int(min=1, max=4))
                ]
            }
        return {}
    
    def create_performance_test_data(self) -> None:
        """Create data for performance testing."""
        perf_dir = self.output_dir / "performance"
        perf_dir.mkdir(exist_ok=True)
        
        # Small document (fast processing)
        self.create_sample_pdf("small_document.pdf", pages=5, with_images=False)
        
        # Medium document (moderate processing)
        self.create_sample_pdf("medium_document.pdf", pages=20, with_images=True)
        
        # Large document (slow processing)
        self.create_sample_pdf("large_document.pdf", pages=50, with_images=True)
        
        # Create search test queries
        search_queries = [
            "machine learning",
            "artificial intelligence",
            "neural networks",
            "data structures",
            "algorithms",
            "deep learning",
            "natural language processing",
            "computer vision",
            "reinforcement learning",
            "supervised learning"
        ]
        
        queries_file = perf_dir / "search_queries.json"
        queries_file.write_text(json.dumps(search_queries, indent=2))
    
    def create_security_test_data(self) -> None:
        """Create data for security testing."""
        security_dir = self.output_dir / "security"
        security_dir.mkdir(exist_ok=True)
        
        # Malicious files
        files = {
            "malicious.exe": b"MZ\x90\x00",  # PE header
            "script.js": b"alert('xss')",
            "payload.php": b"<?php system($_GET['cmd']); ?>",
            "large_file.pdf": b"A" * (100 * 1024 * 1024),  # 100MB
            "empty.pdf": b"",
            "corrupted.pdf": b"invalid pdf content"
        }
        
        for filename, content in files.items():
            file_path = security_dir / filename
            file_path.write_bytes(content)
    
    def setup_all_test_data(self) -> None:
        """Set up all test data."""
        click.echo("🔧 Setting up comprehensive test data...")
        
        # Create sample documents
        click.echo("📄 Creating sample documents...")
        self.create_sample_pdf("sample_small.pdf", pages=5)
        self.create_sample_pdf("sample_medium.pdf", pages=15, with_images=True)
        self.create_sample_pdf("sample_large.pdf", pages=30, with_images=True)
        self.create_sample_docx("sample.docx", pages=8)
        self.create_sample_markdown("sample.md")
        
        # Create database test data
        click.echo("🗄️ Creating database test data...")
        db_data = self.create_test_database_data()
        db_file = self.output_dir / "test_database_data.json"
        db_file.write_text(json.dumps(db_data, indent=2))
        
        # Create performance test data
        click.echo("⚡ Creating performance test data...")
        self.create_performance_test_data()
        
        # Create security test data
        click.echo("🔒 Creating security test data...")
        self.create_security_test_data()
        
        click.echo(f"✅ Test data created in: {self.output_dir}")


@click.command()
@click.option('--output-dir', type=click.Path(), default='test_data', help='Output directory for test data')
@click.option('--clean', is_flag=True, help='Clean existing test data first')
def main(output_dir: str, clean: bool):
    """Set up comprehensive test data for the document learning application."""
    
    output_path = Path(output_dir)
    
    if clean and output_path.exists():
        click.echo(f"🧹 Cleaning existing test data in {output_path}")
        shutil.rmtree(output_path)
    
    generator = TestDataGenerator(output_path)
    generator.setup_all_test_data()


if __name__ == "__main__":
    main()