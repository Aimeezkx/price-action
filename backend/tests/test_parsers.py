"""Comprehensive tests for document parser framework."""

import tempfile
from pathlib import Path
from unittest.mock import Mock, patch, AsyncMock
import asyncio
import pytest
from io import BytesIO

from app.parsers import (
    BaseParser,
    DocxParser,
    ImageData,
    MarkdownParser,
    PDFParser,
    ParsedContent,
    ParserFactory,
    TextBlock,
    get_parser,
    get_parser_for_file,
    get_supported_extensions,
    is_file_supported,
)


class TestBaseParser:
    """Test BaseParser abstract class."""
    
    def test_text_block_creation(self):
        """Test TextBlock creation and validation."""
        text_block = TextBlock(
            text="Sample text",
            page=1,
            bbox={"x": 0, "y": 0, "width": 100, "height": 20},
        )
        
        assert text_block.text == "Sample text"
        assert text_block.page == 1
        assert text_block.bbox["width"] == 100
    
    def test_text_block_invalid_bbox(self):
        """Test TextBlock with invalid bbox."""
        with pytest.raises(ValueError, match="bbox must contain keys"):
            TextBlock(
                text="Sample text",
                page=1,
                bbox={"x": 0, "y": 0},  # Missing width and height
            )
    
    def test_image_data_creation(self):
        """Test ImageData creation and validation."""
        image_data = ImageData(
            image_path="/path/to/image.png",
            page=1,
            bbox={"x": 0, "y": 0, "width": 200, "height": 150},
            format="PNG",
        )
        
        assert image_data.image_path == "/path/to/image.png"
        assert image_data.format == "PNG"
    
    def test_parsed_content_creation(self):
        """Test ParsedContent creation."""
        text_block = TextBlock(
            text="Sample text",
            page=1,
            bbox={"x": 0, "y": 0, "width": 100, "height": 20},
        )
        
        content = ParsedContent(
            text_blocks=[text_block],
            images=[],
            metadata={"format": "test"},
        )
        
        assert len(content.text_blocks) == 1
        assert len(content.images) == 0
        assert content.metadata["format"] == "test"


class MockParser(BaseParser):
    """Mock parser for testing."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = [".mock"]
        self.name = "MockParser"
    
    async def parse(self, file_path: Path) -> ParsedContent:
        """Mock parse implementation."""
        return ParsedContent(
            text_blocks=[
                TextBlock(
                    text="Mock content",
                    page=1,
                    bbox={"x": 0, "y": 0, "width": 100, "height": 20},
                )
            ],
            images=[],
            metadata={"format": "mock"},
        )


class TestParserFactory:
    """Test ParserFactory functionality."""
    
    def test_factory_initialization(self):
        """Test factory initializes with default parsers."""
        factory = ParserFactory()
        
        parsers = factory.list_parsers()
        assert "pdf" in parsers
        assert "docx" in parsers
        assert "markdown" in parsers
    
    def test_register_custom_parser(self):
        """Test registering a custom parser."""
        factory = ParserFactory()
        factory.register_parser("mock", MockParser)
        
        parser = factory.get_parser("mock")
        assert parser is not None
        assert isinstance(parser, MockParser)
    
    def test_register_invalid_parser(self):
        """Test registering invalid parser class."""
        factory = ParserFactory()
        
        with pytest.raises(ValueError, match="must be a subclass of BaseParser"):
            factory.register_parser("invalid", str)
    
    def test_unregister_parser(self):
        """Test unregistering a parser."""
        factory = ParserFactory()
        factory.register_parser("mock", MockParser)
        
        assert factory.get_parser("mock") is not None
        
        factory.unregister_parser("mock")
        assert factory.get_parser("mock") is None
    
    def test_get_parser_for_file(self):
        """Test getting parser for specific file."""
        factory = ParserFactory()
        
        pdf_file = Path("test.pdf")
        parser = factory.get_parser_for_file(pdf_file)
        assert isinstance(parser, PDFParser)
        
        docx_file = Path("test.docx")
        parser = factory.get_parser_for_file(docx_file)
        assert isinstance(parser, DocxParser)
        
        md_file = Path("test.md")
        parser = factory.get_parser_for_file(md_file)
        assert isinstance(parser, MarkdownParser)
        
        unsupported_file = Path("test.xyz")
        parser = factory.get_parser_for_file(unsupported_file)
        assert parser is None
    
    def test_get_supported_extensions(self):
        """Test getting supported extensions."""
        factory = ParserFactory()
        extensions = factory.get_supported_extensions()
        
        assert ".pdf" in extensions
        assert ".docx" in extensions
        assert ".md" in extensions
        assert ".markdown" in extensions
    
    def test_is_file_supported(self):
        """Test checking if file is supported."""
        factory = ParserFactory()
        
        assert factory.is_file_supported(Path("test.pdf"))
        assert factory.is_file_supported(Path("test.docx"))
        assert factory.is_file_supported(Path("test.md"))
        assert not factory.is_file_supported(Path("test.xyz"))


class TestGlobalFunctions:
    """Test global parser functions."""
    
    def test_get_parser(self):
        """Test global get_parser function."""
        parser = get_parser("pdf")
        assert isinstance(parser, PDFParser)
        
        parser = get_parser("nonexistent")
        assert parser is None
    
    def test_get_parser_for_file(self):
        """Test global get_parser_for_file function."""
        parser = get_parser_for_file(Path("test.pdf"))
        assert isinstance(parser, PDFParser)
    
    def test_is_file_supported(self):
        """Test global is_file_supported function."""
        assert is_file_supported(Path("test.pdf"))
        assert not is_file_supported(Path("test.xyz"))
    
    def test_get_supported_extensions(self):
        """Test global get_supported_extensions function."""
        extensions = get_supported_extensions()
        assert ".pdf" in extensions


class TestPDFParser:
    """Test PDFParser functionality."""
    
    def test_pdf_parser_initialization(self):
        """Test PDF parser initialization."""
        parser = PDFParser()
        assert parser.name == "PDFParser"
        assert ".pdf" in parser.supported_extensions
    
    def test_supports_file(self):
        """Test PDF parser file support."""
        parser = PDFParser()
        assert parser.supports_file(Path("test.pdf"))
        assert not parser.supports_file(Path("test.docx"))
    
    def test_validate_file_not_found(self):
        """Test validation with non-existent file."""
        parser = PDFParser()
        
        with pytest.raises(FileNotFoundError):
            parser.validate_file(Path("nonexistent.pdf"))
    
    def test_validate_unsupported_file(self):
        """Test validation with unsupported file."""
        parser = PDFParser()
        
        with tempfile.NamedTemporaryFile(suffix=".txt") as temp_file:
            with pytest.raises(ValueError, match="Unsupported file format"):
                parser.validate_file(Path(temp_file.name))


class TestDocxParser:
    """Test DocxParser functionality."""
    
    def test_docx_parser_initialization(self):
        """Test DOCX parser initialization."""
        parser = DocxParser()
        assert parser.name == "DocxParser"
        assert ".docx" in parser.supported_extensions
        assert ".doc" in parser.supported_extensions
    
    def test_supports_file(self):
        """Test DOCX parser file support."""
        parser = DocxParser()
        assert parser.supports_file(Path("test.docx"))
        assert parser.supports_file(Path("test.doc"))
        assert not parser.supports_file(Path("test.pdf"))


class TestMarkdownParser:
    """Test MarkdownParser functionality."""
    
    def test_markdown_parser_initialization(self):
        """Test Markdown parser initialization."""
        parser = MarkdownParser()
        assert parser.name == "MarkdownParser"
        assert ".md" in parser.supported_extensions
        assert ".markdown" in parser.supported_extensions
    
    def test_supports_file(self):
        """Test Markdown parser file support."""
        parser = MarkdownParser()
        assert parser.supports_file(Path("test.md"))
        assert parser.supports_file(Path("test.markdown"))
        assert not parser.supports_file(Path("test.pdf"))
    
    def test_extract_front_matter(self):
        """Test front matter extraction."""
        parser = MarkdownParser()
        
        content_with_front_matter = """---
title: Test Document
author: Test Author
---

# Main Content

This is the main content."""
        
        content, front_matter = parser._extract_front_matter(content_with_front_matter)
        
        assert "title" in front_matter
        assert front_matter["title"] == "Test Document"
        assert content.startswith("# Main Content")
    
    def test_parse_sections(self):
        """Test markdown section parsing."""
        parser = MarkdownParser()
        
        content = """# Header 1

Some text content.

## Header 2

More content with a list:

- Item 1
- Item 2

```python
code block
```

Final paragraph."""
        
        sections = parser._parse_sections(content)
        
        # Should have headers, text, list, and code sections
        assert len(sections) > 0
        
        # Check for different section types
        section_types = [s["type"] for s in sections]
        assert "header" in section_types
        assert "text" in section_types


@pytest.mark.asyncio
class TestAsyncParsing:
    """Test asynchronous parsing functionality."""
    
    async def test_mock_parser_async(self):
        """Test async parsing with mock parser."""
        parser = MockParser()
        
        with tempfile.NamedTemporaryFile(suffix=".mock") as temp_file:
            temp_path = Path(temp_file.name)
            
            # Mock the file existence check
            with patch.object(Path, "exists", return_value=True):
                result = await parser.parse(temp_path)
                
                assert isinstance(result, ParsedContent)
                assert len(result.text_blocks) == 1
                assert result.text_blocks[0].text == "Mock content"


@pytest.mark.asyncio
class TestPDFParserIntegration:
    """Test PDF parser with real PDF processing."""
    
    async def test_pdf_text_extraction(self, sample_pdf_path):
        """Test PDF text extraction with sample file."""
        parser = PDFParser()
        result = await parser.parse(sample_pdf_path)
        
        assert isinstance(result, ParsedContent)
        assert len(result.text_blocks) > 0
        assert any("Hello World" in block.text for block in result.text_blocks)
        assert all(block.page > 0 for block in result.text_blocks)
    
    async def test_pdf_with_images(self, temp_dir):
        """Test PDF parsing with images."""
        # Create a more complex PDF with images (mock)
        with patch('app.parsers.pdf_parser.fitz') as mock_fitz:
            mock_doc = Mock()
            mock_page = Mock()
            mock_page.get_text.return_value = "Sample text"
            mock_page.get_images.return_value = [
                (0, 0, 100, 100, 8, "DeviceRGB", "", "image1", "DCTDecode")
            ]
            mock_page.get_pixmap.return_value.save.return_value = None
            mock_doc.__iter__ = Mock(return_value=iter([mock_page]))
            mock_doc.page_count = 1
            mock_fitz.open.return_value = mock_doc
            
            parser = PDFParser()
            pdf_path = temp_dir / "test_with_images.pdf"
            pdf_path.write_bytes(b"fake pdf content")
            
            result = await parser.parse(pdf_path)
            
            assert len(result.text_blocks) > 0
            # Images would be extracted in real implementation
    
    async def test_pdf_error_handling(self, temp_dir):
        """Test PDF parser error handling."""
        parser = PDFParser()
        
        # Test with corrupted PDF
        corrupted_pdf = temp_dir / "corrupted.pdf"
        corrupted_pdf.write_bytes(b"not a pdf")
        
        with pytest.raises(Exception):
            await parser.parse(corrupted_pdf)


@pytest.mark.asyncio
class TestDocxParserIntegration:
    """Test DOCX parser with real document processing."""
    
    async def test_docx_text_extraction(self, sample_docx_path):
        """Test DOCX text extraction."""
        parser = DocxParser()
        result = await parser.parse(sample_docx_path)
        
        assert isinstance(result, ParsedContent)
        assert len(result.text_blocks) > 0
        assert any("Test Document" in block.text for block in result.text_blocks)
        assert any("Chapter 1" in block.text for block in result.text_blocks)
    
    async def test_docx_structure_preservation(self, sample_docx_path):
        """Test that DOCX structure is preserved."""
        parser = DocxParser()
        result = await parser.parse(sample_docx_path)
        
        # Check that headings are identified
        headings = [block for block in result.text_blocks if "heading" in block.text.lower()]
        assert len(headings) >= 1
        
        # Check metadata
        assert "format" in result.metadata
        assert result.metadata["format"] == "docx"
    
    async def test_docx_with_images(self, temp_dir):
        """Test DOCX parsing with embedded images."""
        from docx import Document
        from docx.shared import Inches
        
        # Create DOCX with image placeholder
        doc = Document()
        doc.add_heading('Document with Image', 0)
        doc.add_paragraph('Text before image.')
        # In real implementation, would add actual image
        doc.add_paragraph('Text after image.')
        
        docx_path = temp_dir / "with_image.docx"
        doc.save(str(docx_path))
        
        parser = DocxParser()
        result = await parser.parse(docx_path)
        
        assert len(result.text_blocks) >= 3  # Title + 2 paragraphs


@pytest.mark.asyncio
class TestMarkdownParserIntegration:
    """Test Markdown parser with real markdown processing."""
    
    async def test_markdown_parsing(self, sample_markdown_path):
        """Test markdown parsing with sample file."""
        parser = MarkdownParser()
        result = await parser.parse(sample_markdown_path)
        
        assert isinstance(result, ParsedContent)
        assert len(result.text_blocks) > 0
        
        # Check for headers
        headers = [block for block in result.text_blocks 
                  if block.text.startswith('#')]
        assert len(headers) >= 2  # Should have main header and chapter header
    
    async def test_markdown_with_front_matter(self, temp_dir):
        """Test markdown with YAML front matter."""
        md_content = """---
title: Test Document
author: Test Author
tags: [test, markdown]
---

# Main Content

This is the main content of the document.

## Section 1

Some content here.
"""
        
        md_path = temp_dir / "with_frontmatter.md"
        md_path.write_text(md_content)
        
        parser = MarkdownParser()
        result = await parser.parse(md_path)
        
        # Check that front matter is in metadata
        assert "title" in result.metadata
        assert result.metadata["title"] == "Test Document"
        assert "author" in result.metadata
        assert result.metadata["author"] == "Test Author"
        
        # Check that content doesn't include front matter
        content_text = " ".join(block.text for block in result.text_blocks)
        assert "---" not in content_text
        assert "title:" not in content_text
    
    async def test_markdown_code_blocks(self, temp_dir):
        """Test markdown code block handling."""
        md_content = """# Code Examples

Here's some Python code:

```python
def hello_world():
    print("Hello, World!")
    return True
```

And some inline `code` here.

```javascript
console.log("JavaScript example");
```
"""
        
        md_path = temp_dir / "with_code.md"
        md_path.write_text(md_content)
        
        parser = MarkdownParser()
        result = await parser.parse(md_path)
        
        # Check that code blocks are preserved
        content_text = " ".join(block.text for block in result.text_blocks)
        assert "def hello_world" in content_text
        assert "console.log" in content_text
    
    async def test_markdown_lists_and_tables(self, temp_dir):
        """Test markdown list and table parsing."""
        md_content = """# Lists and Tables

## Unordered List
- Item 1
- Item 2
  - Nested item
- Item 3

## Ordered List
1. First item
2. Second item
3. Third item

## Table
| Column 1 | Column 2 | Column 3 |
|----------|----------|----------|
| Row 1    | Data 1   | Value 1  |
| Row 2    | Data 2   | Value 2  |
"""
        
        md_path = temp_dir / "with_lists_tables.md"
        md_path.write_text(md_content)
        
        parser = MarkdownParser()
        result = await parser.parse(md_path)
        
        content_text = " ".join(block.text for block in result.text_blocks)
        assert "Item 1" in content_text
        assert "First item" in content_text
        assert "Column 1" in content_text
        assert "Row 1" in content_text


class TestParserPerformance:
    """Test parser performance characteristics."""
    
    def test_large_document_handling(self, temp_dir):
        """Test parser performance with large documents."""
        # Create a large markdown document
        large_content = "# Large Document\n\n"
        for i in range(1000):
            large_content += f"## Section {i}\n\nThis is section {i} with some content. " * 10
            large_content += "\n\n"
        
        large_md = temp_dir / "large.md"
        large_md.write_text(large_content)
        
        parser = MarkdownParser()
        
        import time
        start_time = time.time()
        
        # Run synchronously for performance test
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            result = loop.run_until_complete(parser.parse(large_md))
            processing_time = time.time() - start_time
            
            # Should process within reasonable time (adjust threshold as needed)
            assert processing_time < 5.0  # 5 seconds max
            assert len(result.text_blocks) > 1000  # Should have many blocks
        finally:
            loop.close()
    
    def test_memory_usage(self, temp_dir):
        """Test parser memory usage with multiple documents."""
        import psutil
        import os
        
        process = psutil.Process(os.getpid())
        initial_memory = process.memory_info().rss
        
        # Parse multiple documents
        for i in range(10):
            md_content = f"# Document {i}\n\n" + "Content " * 1000
            md_path = temp_dir / f"doc_{i}.md"
            md_path.write_text(md_content)
            
            parser = MarkdownParser()
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                result = loop.run_until_complete(parser.parse(md_path))
                assert len(result.text_blocks) > 0
            finally:
                loop.close()
        
        final_memory = process.memory_info().rss
        memory_increase = final_memory - initial_memory
        
        # Memory increase should be reasonable (less than 100MB)
        assert memory_increase < 100 * 1024 * 1024


class TestParserErrorHandling:
    """Test parser error handling and edge cases."""
    
    def test_empty_file_handling(self, temp_dir):
        """Test handling of empty files."""
        empty_pdf = temp_dir / "empty.pdf"
        empty_pdf.write_bytes(b"")
        
        parser = PDFParser()
        
        with pytest.raises(Exception):  # Should raise appropriate error
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                loop.run_until_complete(parser.parse(empty_pdf))
            finally:
                loop.close()
    
    def test_permission_error_handling(self, temp_dir):
        """Test handling of permission errors."""
        restricted_file = temp_dir / "restricted.pdf"
        restricted_file.write_bytes(b"fake pdf")
        
        # Make file unreadable
        restricted_file.chmod(0o000)
        
        parser = PDFParser()
        
        try:
            with pytest.raises((PermissionError, OSError)):
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    loop.run_until_complete(parser.parse(restricted_file))
                finally:
                    loop.close()
        finally:
            # Restore permissions for cleanup
            restricted_file.chmod(0o644)
    
    def test_malformed_content_handling(self, temp_dir):
        """Test handling of malformed content."""
        # Test malformed markdown
        malformed_md = temp_dir / "malformed.md"
        malformed_md.write_text("# Header\n\n```python\n# Unclosed code block")
        
        parser = MarkdownParser()
        
        # Should handle gracefully without crashing
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            result = loop.run_until_complete(parser.parse(malformed_md))
            assert isinstance(result, ParsedContent)
        finally:
            loop.close()


class TestParserConcurrency:
    """Test parser behavior under concurrent access."""
    
    @pytest.mark.asyncio
    async def test_concurrent_parsing(self, temp_dir):
        """Test parsing multiple files concurrently."""
        # Create multiple test files
        files = []
        for i in range(5):
            md_content = f"# Document {i}\n\nContent for document {i}."
            md_path = temp_dir / f"concurrent_{i}.md"
            md_path.write_text(md_content)
            files.append(md_path)
        
        parser = MarkdownParser()
        
        # Parse all files concurrently
        tasks = [parser.parse(file_path) for file_path in files]
        results = await asyncio.gather(*tasks)
        
        assert len(results) == 5
        for i, result in enumerate(results):
            assert isinstance(result, ParsedContent)
            content_text = " ".join(block.text for block in result.text_blocks)
            assert f"Document {i}" in content_text
    
    @pytest.mark.asyncio
    async def test_parser_thread_safety(self, temp_dir):
        """Test parser thread safety."""
        md_content = "# Shared Document\n\nShared content."
        md_path = temp_dir / "shared.md"
        md_path.write_text(md_content)
        
        # Create multiple parser instances
        parsers = [MarkdownParser() for _ in range(3)]
        
        # Parse same file with different parser instances
        tasks = [parser.parse(md_path) for parser in parsers]
        results = await asyncio.gather(*tasks)
        
        assert len(results) == 3
        for result in results:
            assert isinstance(result, ParsedContent)
            content_text = " ".join(block.text for block in result.text_blocks)
            assert "Shared Document" in content_text