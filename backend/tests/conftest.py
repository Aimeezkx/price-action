"""
Pytest configuration and fixtures for comprehensive testing.
"""
import asyncio
import os
import tempfile
from pathlib import Path
from typing import AsyncGenerator, Generator
from unittest.mock import AsyncMock, MagicMock

import pytest
import pytest_asyncio
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.core.config import get_settings
from app.core.database import Base, get_db
from app.main import app
from tests.factories import DocumentFactory, ChapterFactory, KnowledgeFactory, CardFactory


# Test database configuration
TEST_DATABASE_URL = "sqlite+aiosqlite:///:memory:"
TEST_SYNC_DATABASE_URL = "sqlite:///:memory:"

# Create test engines
test_engine = create_async_engine(
    TEST_DATABASE_URL,
    poolclass=StaticPool,
    connect_args={"check_same_thread": False},
)

test_sync_engine = create_engine(
    TEST_SYNC_DATABASE_URL,
    poolclass=StaticPool,
    connect_args={"check_same_thread": False},
)

TestingSessionLocal = sessionmaker(
    test_sync_engine, class_=AsyncSession, expire_on_commit=False
)


@pytest.fixture(scope="session")
def event_loop():
    """Create an instance of the default event loop for the test session."""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest_asyncio.fixture
async def db_session() -> AsyncGenerator[AsyncSession, None]:
    """Create a test database session."""
    async with test_engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
        
        async with AsyncSession(test_engine) as session:
            yield session
            
        await connection.run_sync(Base.metadata.drop_all)


@pytest.fixture
def override_get_db(db_session: AsyncSession):
    """Override the get_db dependency."""
    async def _override_get_db():
        yield db_session
    
    app.dependency_overrides[get_db] = _override_get_db
    yield
    app.dependency_overrides.clear()


@pytest.fixture
def client(override_get_db) -> TestClient:
    """Create a test client."""
    return TestClient(app)


@pytest.fixture
def temp_dir() -> Generator[Path, None, None]:
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as temp_dir:
        yield Path(temp_dir)


@pytest.fixture
def sample_pdf_path(temp_dir: Path) -> Path:
    """Create a sample PDF file for testing."""
    pdf_path = temp_dir / "sample.pdf"
    # Create a minimal PDF content for testing
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj

2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj

3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj

4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
72 720 Td
(Hello World) Tj
ET
endstream
endobj

xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000204 00000 n 
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
297
%%EOF"""
    
    pdf_path.write_bytes(pdf_content)
    return pdf_path


@pytest.fixture
def sample_docx_path(temp_dir: Path) -> Path:
    """Create a sample DOCX file for testing."""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    doc.add_heading('Chapter 1', level=1)
    doc.add_paragraph('This is chapter 1 content.')
    
    docx_path = temp_dir / "sample.docx"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def sample_markdown_path(temp_dir: Path) -> Path:
    """Create a sample Markdown file for testing."""
    markdown_content = """# Test Document

This is a test document.

## Chapter 1

This is chapter 1 content.

### Section 1.1

Some detailed content here.

![Test Image](test_image.png)
"""
    
    md_path = temp_dir / "sample.md"
    md_path.write_text(markdown_content)
    return md_path


@pytest.fixture
def mock_redis():
    """Mock Redis for testing."""
    mock_redis = MagicMock()
    mock_redis.get.return_value = None
    mock_redis.set.return_value = True
    mock_redis.delete.return_value = True
    mock_redis.exists.return_value = False
    return mock_redis


@pytest.fixture
def mock_rq_queue():
    """Mock RQ queue for testing."""
    mock_queue = MagicMock()
    mock_queue.enqueue.return_value = MagicMock(id="test-job-id")
    return mock_queue


@pytest.fixture
def mock_embedding_model():
    """Mock sentence transformer model."""
    mock_model = MagicMock()
    mock_model.encode.return_value = [0.1] * 384  # Mock 384-dim embedding
    return mock_model


@pytest.fixture
def mock_llm_client():
    """Mock LLM client for testing."""
    mock_client = AsyncMock()
    mock_client.generate.return_value = {
        "knowledge_points": [
            {
                "text": "Test knowledge point",
                "kind": "definition",
                "entities": ["test", "knowledge"]
            }
        ]
    }
    return mock_client


# Performance testing fixtures
@pytest.fixture
def performance_test_data():
    """Generate test data for performance testing."""
    return {
        "small_document": "A" * 1000,  # 1KB
        "medium_document": "B" * 100000,  # 100KB
        "large_document": "C" * 1000000,  # 1MB
        "test_queries": [
            "machine learning",
            "artificial intelligence",
            "neural networks",
            "data structures",
            "algorithms"
        ]
    }


# Security testing fixtures
@pytest.fixture
def malicious_files(temp_dir: Path):
    """Create malicious test files."""
    files = {}
    
    # Executable file
    exe_path = temp_dir / "malicious.exe"
    exe_path.write_bytes(b"MZ\x90\x00")  # PE header
    files["exe"] = exe_path
    
    # Script file
    js_path = temp_dir / "script.js"
    js_path.write_text("alert('xss')")
    files["js"] = js_path
    
    # Oversized file
    large_path = temp_dir / "large.pdf"
    large_path.write_bytes(b"A" * (100 * 1024 * 1024))  # 100MB
    files["large"] = large_path
    
    return files


# Factory fixtures
@pytest.fixture
def document_factory():
    """Document factory for creating test documents."""
    return DocumentFactory


@pytest.fixture
def chapter_factory():
    """Chapter factory for creating test chapters."""
    return ChapterFactory


@pytest.fixture
def knowledge_factory():
    """Knowledge factory for creating test knowledge points."""
    return KnowledgeFactory


@pytest.fixture
def card_factory():
    """Card factory for creating test flashcards."""
    return CardFactory


# Test markers
pytest.mark.unit = pytest.mark.unit
pytest.mark.integration = pytest.mark.integration
pytest.mark.performance = pytest.mark.performance
pytest.mark.security = pytest.mark.security
pytest.mark.e2e = pytest.mark.e2e