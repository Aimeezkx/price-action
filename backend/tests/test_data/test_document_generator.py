"""
Test Document Generator

Generates synthetic test documents for various testing scenarios.
"""

import os
import json
import random
import string
from pathlib import Path
from typing import Dict, List, Any
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table
from docx import Document
from docx.shared import Inches
import markdown


class TestDocumentGenerator:
    """Generate synthetic test documents for testing purposes"""
    
    def __init__(self, output_dir: str = "backend/tests/test_data/documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def generate_all_test_documents(self):
        """Generate complete set of test documents"""
        print("Generating test document library...")
        
        # Small documents
        self.generate_small_text_pdf()
        self.generate_small_structured_pdf()
        self.generate_small_docx()
        self.generate_small_markdown()
        
        # Medium documents
        self.generate_medium_textbook_pdf()
        self.generate_medium_report_docx()
        
        # Large documents
        self.generate_large_book_pdf()
        
        # Special format documents
        self.generate_multilingual_pdf()
        self.generate_math_heavy_pdf()
        self.generate_corrupted_pdf()
        
        # Edge cases
        self.generate_empty_pdf()
        self.generate_single_image_pdf()
        
        print("Test document library generation complete!")
        
    def generate_small_text_pdf(self):
        """Generate small text-only PDF (3 pages)"""
        filename = self.output_dir / "small_text.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph("Introduction to Machine Learning", styles['Title']))
        story.append(Spacer(1, 12))
        
        # Content for 3 pages
        content = [
            "Machine learning is a subset of artificial intelligence that focuses on algorithms that can learn from data.",
            "There are three main types of machine learning: supervised, unsupervised, and reinforcement learning.",
            "Supervised learning uses labeled data to train models that can make predictions on new, unseen data.",
            "Common supervised learning algorithms include linear regression, decision trees, and neural networks.",
            "Unsupervised learning finds patterns in data without labeled examples.",
            "Clustering and dimensionality reduction are common unsupervised learning techniques.",
            "Reinforcement learning involves agents learning through interaction with an environment.",
            "Deep learning uses neural networks with multiple layers to learn complex patterns.",
            "Feature engineering is crucial for traditional machine learning success.",
            "Cross-validation helps evaluate model performance and prevent overfitting."
        ]
        
        for i, text in enumerate(content):
            story.append(Paragraph(f"Section {i+1}", styles['Heading2']))
            story.append(Paragraph(text * 3, styles['Normal']))  # Repeat to fill pages
            story.append(Spacer(1, 12))
            
        doc.build(story)
        
    def generate_small_structured_pdf(self):
        """Generate small PDF with clear chapter structure"""
        filename = self.output_dir / "small_structured.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title page
        story.append(Paragraph("Data Structures and Algorithms", styles['Title']))
        story.append(Spacer(1, 24))
        story.append(Paragraph("A Comprehensive Guide", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Table of contents
        story.append(Paragraph("Table of Contents", styles['Heading1']))
        toc_items = [
            "Chapter 1: Arrays and Lists",
            "Chapter 2: Trees and Graphs", 
            "Chapter 3: Sorting Algorithms"
        ]
        for item in toc_items:
            story.append(Paragraph(item, styles['Normal']))
        story.append(Spacer(1, 24))
        
        # Chapters
        chapters = [
            ("Arrays and Lists", "Arrays are fundamental data structures that store elements in contiguous memory locations."),
            ("Trees and Graphs", "Trees are hierarchical data structures with nodes connected by edges."),
            ("Sorting Algorithms", "Sorting algorithms arrange elements in a specific order, such as ascending or descending.")
        ]
        
        for i, (title, content) in enumerate(chapters, 1):
            story.append(Paragraph(f"Chapter {i}: {title}", styles['Heading1']))
            story.append(Paragraph(content * 5, styles['Normal']))
            story.append(Spacer(1, 12))
            
        doc.build(story)
        
    def generate_small_docx(self):
        """Generate small Word document"""
        filename = self.output_dir / "small_document.docx"
        doc = Document()
        
        # Title
        doc.add_heading('Python Programming Basics', 0)
        
        # Introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph('Python is a high-level programming language known for its simplicity and readability.')
        
        # Variables section
        doc.add_heading('Variables and Data Types', level=1)
        doc.add_paragraph('Python supports various data types including integers, floats, strings, and booleans.')
        
        # Code example
        doc.add_paragraph('Example:', style='Intense Quote')
        doc.add_paragraph('x = 10\nname = "Python"\nis_awesome = True')
        
        # Functions section
        doc.add_heading('Functions', level=1)
        doc.add_paragraph('Functions are reusable blocks of code that perform specific tasks.')
        
        doc.save(str(filename))
        
    def generate_small_markdown(self):
        """Generate small Markdown document"""
        filename = self.output_dir / "small_notes.md"
        content = """# Web Development Notes

## HTML Basics

HTML (HyperText Markup Language) is the standard markup language for web pages.

### Key Elements

- `<html>`: Root element
- `<head>`: Document metadata
- `<body>`: Document content
- `<div>`: Generic container
- `<p>`: Paragraph

## CSS Fundamentals

CSS (Cascading Style Sheets) controls the presentation of HTML elements.

### Selectors

1. Element selector: `p { color: blue; }`
2. Class selector: `.highlight { background: yellow; }`
3. ID selector: `#header { font-size: 24px; }`

## JavaScript Basics

JavaScript adds interactivity to web pages.

```javascript
function greet(name) {
    return "Hello, " + name + "!";
}
```

### Variables

- `var`: Function-scoped
- `let`: Block-scoped
- `const`: Block-scoped constant
"""
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(content)
            
    def generate_medium_textbook_pdf(self):
        """Generate medium-sized textbook PDF (15 pages)"""
        filename = self.output_dir / "medium_textbook.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph("Database Systems: Theory and Practice", styles['Title']))
        story.append(Spacer(1, 24))
        
        # Generate multiple chapters
        chapters = [
            ("Introduction to Databases", "Database systems are essential for managing large amounts of structured data."),
            ("Relational Model", "The relational model organizes data into tables with rows and columns."),
            ("SQL Fundamentals", "SQL (Structured Query Language) is used to interact with relational databases."),
            ("Database Design", "Proper database design ensures data integrity and efficient queries."),
            ("Normalization", "Normalization reduces data redundancy and improves data integrity."),
            ("Indexing", "Indexes improve query performance by creating efficient data access paths."),
            ("Transactions", "Transactions ensure data consistency in multi-user environments."),
            ("Concurrency Control", "Concurrency control manages simultaneous database access.")
        ]
        
        for i, (title, intro) in enumerate(chapters, 1):
            story.append(Paragraph(f"Chapter {i}: {title}", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            # Add introduction
            story.append(Paragraph(intro, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Add subsections
            for j in range(3):
                story.append(Paragraph(f"{i}.{j+1} Subsection", styles['Heading2']))
                story.append(Paragraph(f"This subsection covers important aspects of {title.lower()}. " * 10, styles['Normal']))
                story.append(Spacer(1, 12))
                
        doc.build(story)
        
    def generate_medium_report_docx(self):
        """Generate medium-sized business report"""
        filename = self.output_dir / "medium_report.docx"
        doc = Document()
        
        # Title page
        doc.add_heading('Quarterly Business Analysis Report', 0)
        doc.add_paragraph('Q3 2024 Performance Review')
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph('This report analyzes the company performance for Q3 2024, highlighting key metrics and trends.')
        
        # Financial Performance
        doc.add_heading('Financial Performance', level=1)
        doc.add_paragraph('Revenue increased by 15% compared to the previous quarter.')
        
        # Add a table
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Q2 2024'
        hdr_cells[2].text = 'Q3 2024'
        
        # Data rows
        data = [
            ('Revenue', '$2.5M', '$2.9M'),
            ('Profit', '$450K', '$520K'),
            ('Customers', '1,200', '1,380')
        ]
        
        for i, (metric, q2, q3) in enumerate(data, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = metric
            row_cells[1].text = q2
            row_cells[2].text = q3
            
        # Market Analysis
        doc.add_heading('Market Analysis', level=1)
        doc.add_paragraph('The market showed strong growth in our target segments.')
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        doc.add_paragraph('Based on the analysis, we recommend the following actions:')
        doc.add_paragraph('1. Increase marketing spend in high-performing segments', style='List Number')
        doc.add_paragraph('2. Expand product offerings based on customer feedback', style='List Number')
        doc.add_paragraph('3. Optimize operational efficiency to improve margins', style='List Number')
        
        doc.save(str(filename))
        
    def generate_large_book_pdf(self):
        """Generate large book PDF (50+ pages)"""
        filename = self.output_dir / "large_book.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title page
        story.append(Paragraph("Complete Guide to Software Engineering", styles['Title']))
        story.append(Spacer(1, 24))
        story.append(Paragraph("From Fundamentals to Advanced Practices", styles['Heading1']))
        story.append(Spacer(1, 48))
        
        # Generate many chapters to reach 50+ pages
        chapters = [
            "Introduction to Software Engineering",
            "Software Development Life Cycle",
            "Requirements Engineering",
            "System Design and Architecture",
            "Object-Oriented Programming",
            "Design Patterns",
            "Testing Strategies",
            "Version Control Systems",
            "Agile Methodologies",
            "DevOps and Continuous Integration",
            "Database Design",
            "Web Development",
            "Mobile Development",
            "Security Considerations",
            "Performance Optimization",
            "Code Quality and Reviews",
            "Project Management",
            "Team Collaboration",
            "Documentation",
            "Maintenance and Evolution"
        ]
        
        for i, chapter_title in enumerate(chapters, 1):
            story.append(Paragraph(f"Chapter {i}: {chapter_title}", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            # Add multiple sections per chapter
            for j in range(5):
                story.append(Paragraph(f"{i}.{j+1} Section Title", styles['Heading2']))
                
                # Add substantial content
                content = f"This section discusses important aspects of {chapter_title.lower()}. " * 20
                story.append(Paragraph(content, styles['Normal']))
                story.append(Spacer(1, 12))
                
                # Add subsections
                for k in range(3):
                    story.append(Paragraph(f"{i}.{j+1}.{k+1} Subsection", styles['Heading3']))
                    subcontent = f"Detailed information about this subsection. " * 15
                    story.append(Paragraph(subcontent, styles['Normal']))
                    story.append(Spacer(1, 12))
                    
        doc.build(story)
        
    def generate_multilingual_pdf(self):
        """Generate PDF with multiple languages"""
        filename = self.output_dir / "multilingual.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph("Multilingual Document / Document Multilingue / 多言語文書", styles['Title']))
        story.append(Spacer(1, 24))
        
        # English section
        story.append(Paragraph("English Section", styles['Heading1']))
        story.append(Paragraph("This is the English portion of the document. It contains standard English text for testing language detection and processing capabilities.", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # French section
        story.append(Paragraph("Section Française", styles['Heading1']))
        story.append(Paragraph("Ceci est la partie française du document. Elle contient du texte français standard pour tester les capacités de détection et de traitement des langues.", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Spanish section
        story.append(Paragraph("Sección en Español", styles['Heading1']))
        story.append(Paragraph("Esta es la parte en español del documento. Contiene texto estándar en español para probar las capacidades de detección y procesamiento de idiomas.", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # German section
        story.append(Paragraph("Deutscher Abschnitt", styles['Heading1']))
        story.append(Paragraph("Dies ist der deutsche Teil des Dokuments. Er enthält deutschen Standardtext zum Testen der Spracherkennungs- und Verarbeitungsfähigkeiten.", styles['Normal']))
        
        doc.build(story)
        
    def generate_math_heavy_pdf(self):
        """Generate PDF with mathematical content"""
        filename = self.output_dir / "math_heavy.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph("Advanced Calculus and Linear Algebra", styles['Title']))
        story.append(Spacer(1, 24))
        
        # Mathematical content
        story.append(Paragraph("Chapter 1: Differential Calculus", styles['Heading1']))
        story.append(Paragraph("The derivative of a function f(x) is defined as:", styles['Normal']))
        story.append(Paragraph("f'(x) = lim(h→0) [f(x+h) - f(x)] / h", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Common derivatives:", styles['Heading2']))
        derivatives = [
            "d/dx(x^n) = nx^(n-1)",
            "d/dx(e^x) = e^x",
            "d/dx(ln(x)) = 1/x",
            "d/dx(sin(x)) = cos(x)",
            "d/dx(cos(x)) = -sin(x)"
        ]
        
        for deriv in derivatives:
            story.append(Paragraph(f"• {deriv}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Chapter 2: Integration", styles['Heading1']))
        story.append(Paragraph("The fundamental theorem of calculus states:", styles['Normal']))
        story.append(Paragraph("∫[a to b] f'(x)dx = f(b) - f(a)", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Chapter 3: Linear Algebra", styles['Heading1']))
        story.append(Paragraph("Matrix multiplication: C = AB where C[i,j] = Σ(k) A[i,k] * B[k,j]", styles['Normal']))
        story.append(Paragraph("Eigenvalues: Av = λv where λ is the eigenvalue and v is the eigenvector", styles['Normal']))
        
        doc.build(story)
        
    def generate_corrupted_pdf(self):
        """Generate intentionally corrupted PDF for error testing"""
        filename = self.output_dir / "corrupted.pdf"
        
        # Create a normal PDF first
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph("This will be corrupted", styles['Normal'])]
        doc.build(story)
        
        # Corrupt the file by modifying bytes
        with open(filename, 'rb') as f:
            data = bytearray(f.read())
            
        # Corrupt some bytes in the middle
        if len(data) > 100:
            for i in range(50, 100):
                data[i] = random.randint(0, 255)
                
        with open(filename, 'wb') as f:
            f.write(data)
            
    def generate_empty_pdf(self):
        """Generate empty PDF file"""
        filename = self.output_dir / "empty.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        doc.build([])  # Empty story
        
    def generate_single_image_pdf(self):
        """Generate PDF with only one image"""
        filename = self.output_dir / "single_image.pdf"
        
        # Create a simple image first
        from PIL import Image, ImageDraw
        img = Image.new('RGB', (400, 300), color='lightblue')
        draw = ImageDraw.Draw(img)
        draw.text((50, 150), "Test Image", fill='black')
        img_path = self.output_dir / "test_image.png"
        img.save(str(img_path))
        
        # Create PDF with just the image
        doc = SimpleDocTemplate(str(filename), pagesize=letter)
        story = [Image(str(img_path), width=4*Inches, height=3*Inches)]
        doc.build(story)
        
        # Clean up temporary image
        os.remove(str(img_path))


if __name__ == "__main__":
    generator = TestDocumentGenerator()
    generator.generate_all_test_documents()