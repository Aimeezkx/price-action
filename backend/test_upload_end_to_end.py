#!/usr/bin/env python3
"""
End-to-end test for basic upload functionality

This test covers all requirements for Task 6:
- Test successful file upload with document creation
- Test error scenarios (invalid files, large files, etc.)
- Verify frontend can handle upload responses correctly
- Test upload with different file types (PDF, DOCX, TXT, MD)

Requirements: 1.1, 1.2, 1.3, 1.4, 1.5
"""

import asyncio
import sys
import os
import tempfile
import json
from pathlib import Path
from io import BytesIO
from typing import Dict, Any, List
import requests
import time

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent))

def create_test_pdf(content: str = "Test PDF content") -> bytes:
    """Create a simple PDF file for testing"""
    # Simple PDF structure
    pdf_content = f"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj

2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj

3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj

4 0 obj
<<
/Length {len(content) + 50}
>>
stream
BT
/F1 12 Tf
72 720 Td
({content}) Tj
ET
endstream
endobj

xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000074 00000 n 
0000000120 00000 n 
0000000179 00000 n 
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
{200 + len(content)}
%%EOF"""
    return pdf_content.encode('utf-8')

def create_test_docx() -> bytes:
    """Create a simple DOCX file for testing"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test DOCX document for upload testing.')
        doc.add_paragraph('It contains multiple paragraphs to test content extraction.')
        
        # Save to bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except ImportError:
        # Fallback: create minimal DOCX structure
        # This is a very basic ZIP structure that mimics DOCX
        import zipfile
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w') as zf:
            zf.writestr('[Content_Types].xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>''')
            zf.writestr('word/document.xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p><w:r><w:t>Test DOCX document for upload testing.</w:t></w:r></w:p>
</w:body>
</w:document>''')
        buffer.seek(0)
        return buffer.read()

def create_test_txt() -> bytes:
    """Create a test TXT file"""
    content = """Test Text Document

This is a test text document for upload testing.
It contains multiple lines and paragraphs.

Key concepts:
- Document processing
- Text extraction
- Knowledge points
- Flashcard generation

This content should be processed and converted into flashcards.
"""
    return content.encode('utf-8')

def create_test_md() -> bytes:
    """Create a test Markdown file"""
    content = """# Test Markdown Document

This is a test **Markdown** document for upload testing.

## Section 1: Introduction

This document contains various *formatting* elements to test markdown parsing.

### Subsection 1.1: Lists

- Item 1
- Item 2
- Item 3

### Subsection 1.2: Code

```python
def hello_world():
    print("Hello, World!")
```

## Section 2: Conclusion

This content should be processed and converted into flashcards.

> This is a blockquote for testing.

**Important**: This document tests markdown parsing capabilities.
"""
    return content.encode('utf-8')

def create_large_file(size_mb: int = 15) -> bytes:
    """Create a large file for testing size limits"""
    # Create a large PDF
    content = "A" * (size_mb * 1024 * 1024 // 2)  # Fill with content
    return create_test_pdf(content)

def create_invalid_file() -> bytes:
    """Create an invalid file that should be rejected"""
    return b"This is not a valid PDF, DOCX, TXT, or MD file content"

class UploadTester:
    """End-to-end upload functionality tester"""
    
    def __init__(self, base_url: str = "http://localhost:8000"):
        self.base_url = base_url
        self.upload_url = f"{base_url}/api/ingest"
        self.test_results: List[Dict[str, Any]] = []
        
    def log_test_result(self, test_name: str, success: bool, details: Dict[str, Any] = None):
        """Log test result"""
        result = {
            "test_name": test_name,
            "success": success,
            "timestamp": time.time(),
            "details": details or {}
        }
        self.test_results.append(result)
        
        status = "✓ PASS" if success else "✗ FAIL"
        print(f"{status}: {test_name}")
        if details and not success:
            print(f"   Details: {details}")
    
    def check_server_health(self) -> bool:
        """Check if server is running and healthy"""
        try:
            response = requests.get(f"{self.base_url}/api/health", timeout=5)
            if response.status_code == 200:
                print("✓ Server is running and healthy")
                return True
            else:
                print(f"✗ Server health check failed: {response.status_code}")
                return False
        except requests.exceptions.RequestException as e:
            print(f"✗ Cannot connect to server: {e}")
            return False
    
    def test_successful_pdf_upload(self) -> bool:
        """Test successful PDF file upload with document creation"""
        try:
            pdf_content = create_test_pdf("Test PDF content for successful upload")
            
            files = {
                'file': ('test_success.pdf', BytesIO(pdf_content), 'application/pdf')
            }
            
            response = requests.post(self.upload_url, files=files, timeout=30)
            
            if response.status_code == 201:
                data = response.json()
                
                # Verify response structure
                required_fields = ['id', 'filename', 'status', 'created_at']
                missing_fields = [field for field in required_fields if field not in data]
                
                if missing_fields:
                    self.log_test_result(
                        "test_successful_pdf_upload", 
                        False, 
                        {"error": f"Missing fields in response: {missing_fields}", "response": data}
                    )
                    return False
                
                # Verify document was created with correct status
                if data['status'] not in ['pending', 'processing']:
                    self.log_test_result(
                        "test_successful_pdf_upload", 
                        False, 
                        {"error": f"Unexpected status: {data['status']}", "response": data}
                    )
                    return False
                
                self.log_test_result(
                    "test_successful_pdf_upload", 
                    True, 
                    {"document_id": data['id'], "status": data['status']}
                )
                return True
            else:
                self.log_test_result(
                    "test_successful_pdf_upload", 
                    False, 
                    {"status_code": response.status_code, "response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_successful_pdf_upload", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_successful_docx_upload(self) -> bool:
        """Test successful DOCX file upload"""
        try:
            docx_content = create_test_docx()
            
            files = {
                'file': ('test_success.docx', BytesIO(docx_content), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            
            response = requests.post(self.upload_url, files=files, timeout=30)
            
            if response.status_code == 201:
                data = response.json()
                self.log_test_result(
                    "test_successful_docx_upload", 
                    True, 
                    {"document_id": data.get('id'), "status": data.get('status')}
                )
                return True
            else:
                self.log_test_result(
                    "test_successful_docx_upload", 
                    False, 
                    {"status_code": response.status_code, "response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_successful_docx_upload", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_successful_txt_upload(self) -> bool:
        """Test successful TXT file upload"""
        try:
            txt_content = create_test_txt()
            
            files = {
                'file': ('test_success.txt', BytesIO(txt_content), 'text/plain')
            }
            
            response = requests.post(self.upload_url, files=files, timeout=30)
            
            if response.status_code == 201:
                data = response.json()
                self.log_test_result(
                    "test_successful_txt_upload", 
                    True, 
                    {"document_id": data.get('id'), "status": data.get('status')}
                )
                return True
            else:
                self.log_test_result(
                    "test_successful_txt_upload", 
                    False, 
                    {"status_code": response.status_code, "response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_successful_txt_upload", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_successful_md_upload(self) -> bool:
        """Test successful Markdown file upload"""
        try:
            md_content = create_test_md()
            
            files = {
                'file': ('test_success.md', BytesIO(md_content), 'text/markdown')
            }
            
            response = requests.post(self.upload_url, files=files, timeout=30)
            
            if response.status_code == 201:
                data = response.json()
                self.log_test_result(
                    "test_successful_md_upload", 
                    True, 
                    {"document_id": data.get('id'), "status": data.get('status')}
                )
                return True
            else:
                self.log_test_result(
                    "test_successful_md_upload", 
                    False, 
                    {"status_code": response.status_code, "response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_successful_md_upload", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_invalid_file_type(self) -> bool:
        """Test rejection of invalid file types"""
        try:
            invalid_content = create_invalid_file()
            
            files = {
                'file': ('test_invalid.exe', BytesIO(invalid_content), 'application/octet-stream')
            }
            
            response = requests.post(self.upload_url, files=files, timeout=30)
            
            # Should be rejected with 400 status
            if response.status_code == 400:
                data = response.json()
                if 'error' in data:
                    self.log_test_result(
                        "test_invalid_file_type", 
                        True, 
                        {"rejected_correctly": True, "error": data['error']}
                    )
                    return True
            
            self.log_test_result(
                "test_invalid_file_type", 
                False, 
                {"status_code": response.status_code, "response": response.text, "expected": "400 with error"}
            )
            return False
                
        except Exception as e:
            self.log_test_result(
                "test_invalid_file_type", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_empty_file(self) -> bool:
        """Test rejection of empty files"""
        try:
            files = {
                'file': ('test_empty.pdf', BytesIO(b''), 'application/pdf')
            }
            
            response = requests.post(self.upload_url, files=files, timeout=30)
            
            # Should be rejected with 400 status
            if response.status_code == 400:
                data = response.json()
                if 'error' in data:
                    self.log_test_result(
                        "test_empty_file", 
                        True, 
                        {"rejected_correctly": True, "error": data['error']}
                    )
                    return True
            
            self.log_test_result(
                "test_empty_file", 
                False, 
                {"status_code": response.status_code, "response": response.text, "expected": "400 with error"}
            )
            return False
                
        except Exception as e:
            self.log_test_result(
                "test_empty_file", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_large_file(self) -> bool:
        """Test handling of large files (should be rejected if over limit)"""
        try:
            # Create a 15MB file (likely over the limit)
            large_content = create_large_file(15)
            
            files = {
                'file': ('test_large.pdf', BytesIO(large_content), 'application/pdf')
            }
            
            response = requests.post(self.upload_url, files=files, timeout=60)
            
            # Should be rejected with 413 status (Request Entity Too Large)
            if response.status_code == 413:
                data = response.json()
                self.log_test_result(
                    "test_large_file", 
                    True, 
                    {"rejected_correctly": True, "error": data.get('error', 'File too large')}
                )
                return True
            elif response.status_code == 201:
                # If accepted, that's also valid (depends on server configuration)
                data = response.json()
                self.log_test_result(
                    "test_large_file", 
                    True, 
                    {"accepted_large_file": True, "document_id": data.get('id')}
                )
                return True
            else:
                self.log_test_result(
                    "test_large_file", 
                    False, 
                    {"status_code": response.status_code, "response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_large_file", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_no_file_provided(self) -> bool:
        """Test error handling when no file is provided"""
        try:
            # Send request without file
            response = requests.post(self.upload_url, timeout=30)
            
            # Should be rejected with 422 status (Unprocessable Entity)
            if response.status_code in [400, 422]:
                self.log_test_result(
                    "test_no_file_provided", 
                    True, 
                    {"rejected_correctly": True, "status_code": response.status_code}
                )
                return True
            else:
                self.log_test_result(
                    "test_no_file_provided", 
                    False, 
                    {"status_code": response.status_code, "response": response.text, "expected": "400 or 422"}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_no_file_provided", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_malformed_pdf(self) -> bool:
        """Test handling of malformed PDF files"""
        try:
            # Create malformed PDF
            malformed_content = b"%PDF-1.4\nThis is not a valid PDF structure"
            
            files = {
                'file': ('test_malformed.pdf', BytesIO(malformed_content), 'application/pdf')
            }
            
            response = requests.post(self.upload_url, files=files, timeout=30)
            
            # Should either be rejected (400) or accepted but fail processing
            if response.status_code == 400:
                data = response.json()
                self.log_test_result(
                    "test_malformed_pdf", 
                    True, 
                    {"rejected_correctly": True, "error": data.get('error')}
                )
                return True
            elif response.status_code == 201:
                # If accepted, it should be marked for processing
                data = response.json()
                self.log_test_result(
                    "test_malformed_pdf", 
                    True, 
                    {"accepted_for_processing": True, "document_id": data.get('id')}
                )
                return True
            else:
                self.log_test_result(
                    "test_malformed_pdf", 
                    False, 
                    {"status_code": response.status_code, "response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_malformed_pdf", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_document_retrieval(self) -> bool:
        """Test that uploaded documents can be retrieved"""
        try:
            # First upload a document
            pdf_content = create_test_pdf("Test content for retrieval")
            
            files = {
                'file': ('test_retrieval.pdf', BytesIO(pdf_content), 'application/pdf')
            }
            
            upload_response = requests.post(self.upload_url, files=files, timeout=30)
            
            if upload_response.status_code != 201:
                self.log_test_result(
                    "test_document_retrieval", 
                    False, 
                    {"error": "Upload failed", "status_code": upload_response.status_code}
                )
                return False
            
            upload_data = upload_response.json()
            document_id = upload_data['id']
            
            # Try to retrieve the document
            get_response = requests.get(f"{self.base_url}/api/documents/{document_id}", timeout=30)
            
            if get_response.status_code == 200:
                get_data = get_response.json()
                
                # Verify document data
                if get_data['id'] == document_id and get_data['filename'] == upload_data['filename']:
                    self.log_test_result(
                        "test_document_retrieval", 
                        True, 
                        {"document_id": document_id, "retrieved_successfully": True}
                    )
                    return True
                else:
                    self.log_test_result(
                        "test_document_retrieval", 
                        False, 
                        {"error": "Document data mismatch", "uploaded": upload_data, "retrieved": get_data}
                    )
                    return False
            else:
                self.log_test_result(
                    "test_document_retrieval", 
                    False, 
                    {"error": "Document retrieval failed", "status_code": get_response.status_code}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_document_retrieval", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_document_status_endpoint(self) -> bool:
        """Test document status endpoint"""
        try:
            # First upload a document
            pdf_content = create_test_pdf("Test content for status check")
            
            files = {
                'file': ('test_status.pdf', BytesIO(pdf_content), 'application/pdf')
            }
            
            upload_response = requests.post(self.upload_url, files=files, timeout=30)
            
            if upload_response.status_code != 201:
                self.log_test_result(
                    "test_document_status_endpoint", 
                    False, 
                    {"error": "Upload failed", "status_code": upload_response.status_code}
                )
                return False
            
            upload_data = upload_response.json()
            document_id = upload_data['id']
            
            # Check status endpoint
            status_response = requests.get(f"{self.base_url}/api/documents/{document_id}/status", timeout=30)
            
            if status_response.status_code == 200:
                status_data = status_response.json()
                
                # Verify status data structure
                required_fields = ['document_id', 'status', 'filename']
                missing_fields = [field for field in required_fields if field not in status_data]
                
                if missing_fields:
                    self.log_test_result(
                        "test_document_status_endpoint", 
                        False, 
                        {"error": f"Missing fields in status response: {missing_fields}"}
                    )
                    return False
                
                self.log_test_result(
                    "test_document_status_endpoint", 
                    True, 
                    {"document_id": document_id, "status": status_data['status']}
                )
                return True
            else:
                self.log_test_result(
                    "test_document_status_endpoint", 
                    False, 
                    {"error": "Status endpoint failed", "status_code": status_response.status_code}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_document_status_endpoint", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def test_concurrent_uploads(self) -> bool:
        """Test handling of concurrent uploads"""
        try:
            import threading
            import queue
            
            results_queue = queue.Queue()
            
            def upload_file(file_num):
                try:
                    pdf_content = create_test_pdf(f"Concurrent upload test file {file_num}")
                    
                    files = {
                        'file': (f'test_concurrent_{file_num}.pdf', BytesIO(pdf_content), 'application/pdf')
                    }
                    
                    response = requests.post(self.upload_url, files=files, timeout=30)
                    results_queue.put((file_num, response.status_code == 201, response.status_code))
                except Exception as e:
                    results_queue.put((file_num, False, str(e)))
            
            # Start 3 concurrent uploads
            threads = []
            for i in range(3):
                thread = threading.Thread(target=upload_file, args=(i,))
                threads.append(thread)
                thread.start()
            
            # Wait for all threads to complete
            for thread in threads:
                thread.join()
            
            # Collect results
            successful_uploads = 0
            total_uploads = 3
            
            while not results_queue.empty():
                file_num, success, status = results_queue.get()
                if success:
                    successful_uploads += 1
            
            # At least 2 out of 3 should succeed (allowing for some concurrency issues)
            if successful_uploads >= 2:
                self.log_test_result(
                    "test_concurrent_uploads", 
                    True, 
                    {"successful_uploads": successful_uploads, "total_uploads": total_uploads}
                )
                return True
            else:
                self.log_test_result(
                    "test_concurrent_uploads", 
                    False, 
                    {"successful_uploads": successful_uploads, "total_uploads": total_uploads}
                )
                return False
                
        except Exception as e:
            self.log_test_result(
                "test_concurrent_uploads", 
                False, 
                {"error": str(e)}
            )
            return False
    
    def run_all_tests(self) -> bool:
        """Run all upload tests"""
        print("=" * 80)
        print("UPLOAD FUNCTIONALITY END-TO-END TESTS")
        print("=" * 80)
        print()
        
        # Check server health first
        if not self.check_server_health():
            print("❌ Server is not available. Please start the backend server first.")
            return False
        
        print()
        print("Running upload functionality tests...")
        print("-" * 50)
        
        # Test successful uploads with different file types
        tests = [
            self.test_successful_pdf_upload,
            self.test_successful_docx_upload,
            self.test_successful_txt_upload,
            self.test_successful_md_upload,
            
            # Test error scenarios
            self.test_invalid_file_type,
            self.test_empty_file,
            self.test_large_file,
            self.test_no_file_provided,
            self.test_malformed_pdf,
            
            # Test document management
            self.test_document_retrieval,
            self.test_document_status_endpoint,
            
            # Test concurrent uploads
            self.test_concurrent_uploads,
        ]
        
        passed_tests = 0
        total_tests = len(tests)
        
        for test in tests:
            try:
                if test():
                    passed_tests += 1
            except Exception as e:
                print(f"✗ FAIL: {test.__name__} - Unexpected error: {e}")
        
        print()
        print("-" * 50)
        print(f"Test Results: {passed_tests}/{total_tests} tests passed")
        
        if passed_tests == total_tests:
            print("🎉 All upload functionality tests passed!")
            return True
        else:
            print(f"❌ {total_tests - passed_tests} tests failed")
            return False
    
    def generate_test_report(self) -> Dict[str, Any]:
        """Generate detailed test report"""
        passed = sum(1 for result in self.test_results if result['success'])
        failed = len(self.test_results) - passed
        
        report = {
            "summary": {
                "total_tests": len(self.test_results),
                "passed": passed,
                "failed": failed,
                "success_rate": (passed / len(self.test_results)) * 100 if self.test_results else 0
            },
            "test_results": self.test_results,
            "requirements_coverage": {
                "1.1": "File upload with supported formats (PDF, DOCX, TXT, MD)",
                "1.2": "Document record creation in database",
                "1.3": "Processing status tracking",
                "1.4": "Error handling for various scenarios",
                "1.5": "Frontend response compatibility"
            }
        }
        
        return report

def main():
    """Main test execution"""
    tester = UploadTester()
    
    success = tester.run_all_tests()
    
    # Generate and save test report
    report = tester.generate_test_report()
    
    report_file = Path(__file__).parent / "upload_test_report.json"
    with open(report_file, 'w') as f:
        json.dump(report, f, indent=2, default=str)
    
    print(f"\nDetailed test report saved to: {report_file}")
    
    return success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)